import docx
from docx.shared import Pt

doc=docx.Document()

for i in range(2):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run('Thomas is a Clown'+str(i))
    font = run.font
    if i==0:
        font.name = 'Calibri'
    if i==1:
        font.name = 'Arial Black'    
    font.size = Pt(12)

doc.save('new.docx')

