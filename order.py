import io
import random
import docx
from docx.shared import Pt

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

f = io.open("law3.txt", mode="r", encoding="utf-8")
full=f.read()
f.close()


doc=docx.Document()
lsttxt="\n"
#cv2.imshow("test",wincap.windowcap())
last=0
for i in range(len(full)):
    print(i)
    txt=full[i]
    if txt=="J":
        txt="j"
    if txt=="U":
        txt="u"
    check=0
    if txt=="Ü":
        txt="ü"
        check=1
    if txt=="I":
        txt="ı"
        check=1
    if txt=="İ":
        txt="i"
        check=1
    if txt=="Ğ":
        txt="ğ"
        check=1
    if txt=="Ö":
        txt="ö"
        check=1
    if check==0:
        txt=txt.lower()
        """
    """
    rnd=random.randrange(0,3)
    while rnd==last:
        rnd=random.randrange(0,3)
    if txt!=" ":
        last=rnd
    if lsttxt=="\n":
        rn=random.randrange(0,3)
        x=" "*rn
        txt=x+txt
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(txt)
    font = run.font
    if rnd==0:
        font.name = 'yagifont1'
    if rnd==1:
        font.name = 'yagifont2'  
    if rnd==2:
        font.name = 'yagifont3'
    """  
    if rnd==0:
        font.name = 'Doruk1'
    if rnd==1:
        font.name = 'Doruk2'  
    if rnd==2:
        font.name = 'Doruk3'
    """   
    font.size = Pt(12)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("xx")
    lsttxt=txt
    """
for i in doc.paragraphs:
    if i.text=="^pmrk^p":
        delete_paragraph(i)"""
doc.save('yagi3.docx')
    




